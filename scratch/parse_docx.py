try:
    import docx
    doc = docx.Document("Mô tả phân tích nhóm kinh doanh.docx")
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text.append(" | ".join(row_text))
    
    with open("scratch/docx_content.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(text))
    print("Done docx")
except Exception as e:
    import traceback
    print("Error:", traceback.format_exc())
